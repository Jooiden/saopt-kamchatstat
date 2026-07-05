from docx import Document
from docx.shared import Inches
import datetime
import os
import pandas as pd # Для работы с Excel

def generate_word_report(user_name, test_results, chart_path, interpretation_text=""):
    output_dir = "Отчеты_САОПТ"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    doc = Document()
    # Делаем заголовок более официальным
    doc.add_heading('ИНДИВИДУАЛЬНЫЙ ОТЧЕТ ПО РЕЗУЛЬТАТАМ ТЕСТИРОВАНИЯ', 0)

    now = datetime.datetime.now()
    timestamp = now.strftime("%Y-%m-%d_%H-%M")
    
    # Блок информации
    p = doc.add_paragraph()
    p.add_run('Сотрудник: ').bold = True
    p.add_run(user_name)
    p.add_run(f'\nДата формирования: {now.strftime("%d.%m.%Y")}')

    # --- Психологический комментарий (Заключение) ---
    if interpretation_text:
        doc.add_heading('Психологическое заключение', level=1)
        # Убираем из текста слова "ВЫСОКИЙ РИСК", если они вдруг проскочили
        clean_text = interpretation_text.replace("ВЫСОКИЙ РИСК", "Повышенный показатель")
        doc.add_paragraph(clean_text)

    # Таблица результатов
    doc.add_heading('Сводные данные', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Дата'
    hdr_cells[1].text = 'Методика'
    hdr_cells[2].text = 'Результат'

    for res in test_results:
        row_cells = table.add_row().cells
        row_cells[0].text = str(res[2])[:10]
        row_cells[1].text = str(res[0])
        # Вместо "ст." пишем "стн." (станайны) - это грамотнее
        row_cells[2].text = f"{res[1]} стн."

    # График (если есть)
    if os.path.exists(chart_path):
        doc.add_heading('Визуализация показателей', level=1)
        doc.add_picture(chart_path, width=Inches(5.0))

    # Подпись
    doc.add_paragraph('\n\n____________________ / Специалист по кадрам /')

    safe_name = user_name.replace(" ", "_")
    filename = f"Отчет_{safe_name}_{timestamp}.docx"
    full_path = os.path.join(output_dir, filename)
    doc.save(full_path)
    
    return full_path

def generate_group_excel(data):
    import pandas as pd
    import os
    from datetime import datetime
    
    # 1. Формируем дату для названия файла
    # strftime("%Y-%m-%d_%H-%M") формат ГГГГ-ММ-ДД_ЧЧ-ММ
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
    filename = f"Групповой_отчет_{timestamp}.xlsx"
    
    # 2. Подготовка данных 
    excel_data = [row[:4] for row in data]
    df = pd.DataFrame(excel_data, columns=["Сотрудник", "Методика", "Результат (стены)", "Дата прохождения"])
    
    # 3. Создание Excel с оформлением
    writer = pd.ExcelWriter(filename, engine='xlsxwriter')
    df.to_excel(writer, index=False, sheet_name='Результаты')
    
    workbook  = writer.book
    worksheet = writer.sheets['Результаты']
    
    # Стилизация
    
    # Стиль для заголовка
    header_fmt = workbook.add_format({
        'bold': True, 
        'fg_color': '#4F81BD', 
        'font_color': 'white', 
        'border': 1,
        'align': 'center'
    })
    
    # Стиль для обычных ячеек
    cell_fmt = workbook.add_format({'border': 1, 'align': 'center'})
    
    # Стиль для критических значений
    danger_fmt = workbook.add_format({
        'bg_color': '#FFC7CE',
        'font_color': '#9C0006',
        'bold': True,
        'border': 1
    })

    for col_num, value in enumerate(df.columns.values):
        worksheet.write(0, col_num, value, header_fmt)
    
    # Автоподбор ширины столбцов
    for i, col in enumerate(df.columns):
        column_len = max(df[col].astype(str).str.len().max(), len(col)) + 5
        worksheet.set_column(i, i, column_len, cell_fmt)
        
    # форматирование "Результат"
    # Критический балл
    worksheet.conditional_format(1, 2, len(data), 2, {
        'type':     'cell',
        'criteria': '>=',
        'value':    8,
        'format':   danger_fmt
    })

    writer.close()
    return filename